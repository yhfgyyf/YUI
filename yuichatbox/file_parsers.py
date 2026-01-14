"""
File parsing utilities for YUI ChatBox
Extracts text content from various file types
"""

import io
import os
from typing import Optional, Dict, Any
from pathlib import Path

# PDF
try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

# Word
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Images
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# File type detection
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False


class FileParseResult:
    """Result of file parsing operation"""
    def __init__(
        self,
        success: bool,
        text: str = "",
        metadata: Optional[Dict[str, Any]] = None,
        error: Optional[str] = None,
        truncated: bool = False
    ):
        self.success = success
        self.text = text
        self.metadata = metadata or {}
        self.error = error
        self.truncated = truncated


MAX_TEXT_LENGTH = 50000  # Maximum characters to extract


def detect_file_type(file_path: str) -> str:
    """Detect file type using magic numbers"""
    if HAS_MAGIC:
        try:
            mime = magic.from_file(file_path, mime=True)
            return mime
        except Exception:
            pass

    # Fallback to extension
    ext = Path(file_path).suffix.lower()
    mime_map = {
        '.txt': 'text/plain',
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
    }
    return mime_map.get(ext, 'application/octet-stream')


def parse_text_file(file_path: str) -> FileParseResult:
    """Parse plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read(MAX_TEXT_LENGTH + 1)
            truncated = len(text) > MAX_TEXT_LENGTH
            if truncated:
                text = text[:MAX_TEXT_LENGTH]

        return FileParseResult(
            success=True,
            text=text,
            metadata={'encoding': 'utf-8'},
            truncated=truncated
        )
    except Exception as e:
        return FileParseResult(success=False, error=str(e))


def parse_pdf(file_path: str) -> FileParseResult:
    """Parse PDF file"""
    if not HAS_PDF:
        return FileParseResult(success=False, error="PyPDF2 not installed")

    try:
        reader = PdfReader(file_path)
        pages = []
        total_chars = 0
        truncated = False

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if total_chars + len(text) > MAX_TEXT_LENGTH:
                remaining = MAX_TEXT_LENGTH - total_chars
                pages.append(f"[Page {page_num + 1}]\n{text[:remaining]}")
                truncated = True
                break
            pages.append(f"[Page {page_num + 1}]\n{text}")
            total_chars += len(text)

        full_text = "\n\n".join(pages)
        metadata = {
            'pages': len(reader.pages),
            'extracted_pages': len(pages)
        }

        # Add PDF metadata if available
        if reader.metadata:
            metadata.update({
                'title': reader.metadata.get('/Title', ''),
                'author': reader.metadata.get('/Author', ''),
            })

        return FileParseResult(
            success=True,
            text=full_text,
            metadata=metadata,
            truncated=truncated
        )
    except Exception as e:
        return FileParseResult(success=False, error=str(e))


def parse_docx(file_path: str) -> FileParseResult:
    """Parse Word document (.docx)"""
    if not HAS_DOCX:
        return FileParseResult(success=False, error="python-docx not installed")

    try:
        doc = Document(file_path)
        paragraphs = []
        total_chars = 0
        truncated = False

        for para in doc.paragraphs:
            text = para.text
            if total_chars + len(text) > MAX_TEXT_LENGTH:
                remaining = MAX_TEXT_LENGTH - total_chars
                paragraphs.append(text[:remaining])
                truncated = True
                break
            paragraphs.append(text)
            total_chars += len(text) + 1  # +1 for newline

        full_text = "\n".join(paragraphs)
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'extracted_paragraphs': len(paragraphs)
        }

        # Add core properties if available
        if doc.core_properties:
            metadata.update({
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
            })

        return FileParseResult(
            success=True,
            text=full_text,
            metadata=metadata,
            truncated=truncated
        )
    except Exception as e:
        return FileParseResult(success=False, error=str(e))


def parse_image(file_path: str) -> FileParseResult:
    """Parse image file - extract EXIF and basic metadata (no OCR)"""
    if not HAS_PIL:
        return FileParseResult(success=False, error="Pillow not installed")

    try:
        img = Image.open(file_path)
        metadata = {
            'format': img.format,
            'size': img.size,
            'mode': img.mode,
        }

        # Extract EXIF data
        exif_data = img.getexif()
        if exif_data:
            exif_dict = {
                'ExifImageWidth': exif_data.get(0xA002),
                'ExifImageHeight': exif_data.get(0xA003),
                'DateTime': exif_data.get(0x0132),
                'Make': exif_data.get(0x010F),
                'Model': exif_data.get(0x0110),
            }
            metadata['exif'] = {k: v for k, v in exif_dict.items() if v}

        text_parts = [
            f"Image: {Path(file_path).name}",
            f"Format: {img.format}, Size: {img.size[0]}x{img.size[1]}, Mode: {img.mode}"
        ]

        # Add EXIF info if available
        if 'exif' in metadata and metadata['exif']:
            text_parts.append(f"EXIF Data: {metadata['exif']}")

        full_text = "\n".join(text_parts)

        return FileParseResult(
            success=True,
            text=full_text,
            metadata=metadata,
            truncated=False
        )
    except Exception as e:
        return FileParseResult(success=False, error=str(e))


def parse_file(file_path: str) -> FileParseResult:
    """
    Parse file based on its type

    Args:
        file_path: Path to the file

    Returns:
        FileParseResult object
    """
    mime_type = detect_file_type(file_path)

    # Route to appropriate parser
    if mime_type.startswith('text/'):
        return parse_text_file(file_path)
    elif mime_type == 'application/pdf':
        return parse_pdf(file_path)
    elif mime_type in [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    ]:
        return parse_docx(file_path)
    elif mime_type.startswith('image/'):
        return parse_image(file_path)
    else:
        return FileParseResult(
            success=False,
            error=f"Unsupported file type: {mime_type}"
        )
